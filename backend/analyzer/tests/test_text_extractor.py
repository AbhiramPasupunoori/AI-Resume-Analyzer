import tempfile
from pathlib import Path

import pymupdf
from django.test import SimpleTestCase
from docx import Document

from analyzer.services.text_extractor import extract_docx_text, extract_pdf_text


class TextExtractorTests(SimpleTestCase):
    def test_extracts_pdf_text(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "resume.pdf"
            document = pymupdf.open()
            page = document.new_page()
            page.insert_text((72, 72), "Python developer with Django experience")
            document.save(path)
            document.close()

            result = extract_pdf_text(path)

        self.assertIn("Python developer", result)

    def test_extracts_docx_paragraph_and_table_text(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "resume.docx"
            document = Document()
            document.add_paragraph("Backend developer")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Python"
            table.cell(0, 1).text = "Django"
            document.save(path)

            result = extract_docx_text(path)

        self.assertIn("Backend developer", result)
        self.assertIn("Python | Django", result)
